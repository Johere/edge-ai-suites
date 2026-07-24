"""
Report Template Manager.

Handles loading .docx templates, extracting their structure (sections + placeholders),
and filling templates with LLM-generated content.

Template format:
- Headings define sections
- Text inside {placeholder} marks fields the LLM should fill
- Static text (without placeholders) is preserved as-is
"""

import os
import re
import json
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utils.runtime_config_loader import RuntimeConfig

logger = logging.getLogger(__name__)

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")


def _safe_template_path(name: str) -> str:
    """Resolve a template name to a path inside TEMPLATES_DIR.

    Guards against path traversal (only the basename is honored) and forces a
    .docx extension so callers can pass names with or without it.
    """
    if not name:
        raise ValueError("Template name is required.")
    base = os.path.basename(name.strip())
    if not base or base in (".", ".."):
        raise ValueError(f"Invalid template name: {name!r}")
    if not base.lower().endswith(".docx"):
        base += ".docx"
    return os.path.join(TEMPLATES_DIR, base)


def get_template_path(language: str = "zh", session_id: str = None, template_name: str = None) -> str:
    """Get the active template path.

    Resolution order:
      1. An explicitly selected ``template_name`` from the unified library.
      2. A session-level custom template.
      3. A project-level custom template.
      4. The built-in default for the language.
    """
    if template_name:
        selected = _safe_template_path(template_name)
        if os.path.exists(selected):
            return selected
        logger.warning(f"Selected template {template_name!r} not found; falling back to defaults.")

    if session_id:
        project_config = RuntimeConfig.get_section("Project")
        custom_path = os.path.join(
            project_config.get("location"),
            project_config.get("name"),
            session_id,
            "custom_report_template.docx",
        )
        if os.path.exists(custom_path):
            return custom_path

    project_config = RuntimeConfig.get_section("Project")
    project_custom = os.path.join(
        project_config.get("location"),
        project_config.get("name"),
        "report_template.docx",
    )
    if os.path.exists(project_custom):
        return project_custom

    default_path = os.path.join(TEMPLATES_DIR, f"report_template_{language}.docx")
    if os.path.exists(default_path):
        return default_path

    return None


def read_docx_as_markdown(docx_path: str, session_id: str) -> str:
    """Read a filled .docx and convert to markdown with heading structure.

    Embedded images (e.g. the rendered mind map) are emitted as markdown images
    pointing at the session-scoped report image endpoint.
    """
    doc = Document(docx_path)
    lines = []
    list_counter = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', '').replace('Heading', '1'))
                lines.append(f"\n{'#' * level} {text}\n")
                list_counter = 0
            elif para.style.name == 'Title':
                lines.append(f"\n# {text}\n")
                list_counter = 0
            elif para.style.name.startswith('List'):
                list_counter += 1
                normalized = re.sub(r'^\s*\d+\s*[\.、]\s*', '', text).strip()
                lines.append(f"{list_counter}. {normalized}")
            else:
                lines.append(text)
        img_md = _paragraph_image_markdown(doc, para, session_id)
        if img_md:
            lines.append(img_md)
    return "\n\n".join(lines)


def _paragraph_image_markdown(doc, paragraph, session_id: str) -> str:
    """If a paragraph embeds an image, return it as a markdown image URL;
    otherwise return ''.
    """
    blips = paragraph._p.findall('.//' + qn('a:blip'))
    if not blips:
        return ""
    rid = blips[0].get(qn('r:embed'))
    if not rid:
        return ""
    if doc.part.related_parts.get(rid) is None:
        return ""
    return f"\n![Mind map](/report/{session_id}/mindmap-image)\n"


def extract_template_structure(template_path: str) -> dict:
    """Extract the structure from a .docx template.

    Returns a dict with:
      - sections: list of {heading, level, fields: [field_names]}
      - all_fields: flat list of all placeholder names
      - raw_text: full template text for LLM reference
    """
    doc = Document(template_path)
    sections = []
    all_fields = []
    raw_lines = []
    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        raw_lines.append(text)

        if para.style.name.startswith('Heading'):
            level = int(para.style.name.replace('Heading ', '').replace('Heading', '1'))
            current_section = {"heading": text, "level": level, "fields": []}
            sections.append(current_section)
        else:
            placeholders = re.findall(r'\{(\w+)\}', text)
            if placeholders:
                all_fields.extend(placeholders)
                if current_section:
                    current_section["fields"].extend(placeholders)

    return {
        "sections": sections,
        "all_fields": list(dict.fromkeys(all_fields)),
        "raw_text": "\n".join(raw_lines),
    }


def _paragraph_placeholders(paragraph) -> list:
    """The ``{field}`` placeholder codes present in a paragraph (in order)."""
    return re.findall(r'\{(\w+)\}', paragraph.text)


def _remove_paragraph(paragraph) -> None:
    """Detach a paragraph from its parent so it disappears from the document."""
    element = paragraph._element
    element.getparent().remove(element)


def _is_heading(paragraph) -> bool:
    style = paragraph.style.name if paragraph.style else ""
    return style.startswith("Heading") or style == "Title"


def fill_template(template_path: str, field_values: dict, output_path: str,
                  drop_codes: frozenset = frozenset(), image_fields: dict = None,
                  field_codes: frozenset = frozenset()) -> str:
    """Fill a .docx template with values and save to output_path.

    ``drop_codes`` are field codes the user deselected: any paragraph whose
    placeholders are ALL in ``drop_codes`` is removed entirely (so an unchecked
    field's line vanishes rather than rendering an empty label). A section
    ``Heading`` left with no surviving body paragraphs is dropped too, so no
    orphan headings remain.

    ``image_fields`` maps a placeholder code to an image file path: a paragraph
    holding such a placeholder gets the picture inserted in place of the text
    (used for the rendered mind map). Deselection still applies via ``drop_codes``.
    """
    image_fields = image_fields or {}
    doc = Document(template_path)

    if drop_codes:
        field_values = {**field_values, **{c: "" for c in drop_codes}}
        _drop_deselected_paragraphs(doc, drop_codes, field_codes)
        _renumber_sections(doc)

    _number_list_items(doc)

    for para in doc.paragraphs:
        img_code = next((c for c in _paragraph_placeholders(para) if c in image_fields), None)
        if img_code:
            _insert_image_in_paragraph(para, img_code, image_fields[img_code])
        else:
            _replace_placeholders_in_paragraph(para, field_values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_placeholders_in_paragraph(para, field_values)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"Template-based report saved to {output_path}")
    return output_path


def _insert_image_in_paragraph(paragraph, code: str, image_path: str,
                               width_inches: float = 6.0) -> None:
    """Insert an embedded image where a ``{code}`` placeholder sits, scaled to
    ``width_inches`` (fit within the page).

    Any surrounding label text (e.g. "Mind-map view:") is preserved and the image
    starts on a new line below it. Because the label shares the placeholder's
    paragraph, deselecting the field drops the label with it (no orphan). Falls
    back to just the label text if the image can't be added.
    """
    remaining = paragraph.text.replace('{' + code + '}', '').strip()
    for i, run in enumerate(paragraph.runs):
        run.text = remaining if i == 0 else ""
    run = paragraph.add_run()
    try:
        if remaining:
            run.add_break()
        run.add_picture(image_path, width=Inches(width_inches))
    except Exception as e:  # bad/missing image file — don't break the whole report
        logger.warning("Could not embed image %s: %s", image_path, e)


def _drop_deselected_paragraphs(doc, drop_codes: frozenset,
                                field_codes: frozenset = frozenset()) -> None:
    """Remove paragraphs for deselected fields, then any now-empty section heading.

    A paragraph is removed when it has ≥1 placeholder and every placeholder is in
    ``drop_codes``. A ``Heading`` is removed when its section (paragraphs until the
    next heading) has no surviving *field* content.

    ``field_codes`` is the set of real report field codes. Section "body" counts
    only a paragraph holding a surviving field placeholder (one in ``field_codes``
    and not in ``drop_codes``) — so a document footer that trails the last heading
    (a divider / report-notes / confidence line with no catalog field) does NOT
    keep an otherwise-empty heading alive. Falls back to "any non-blank text" when
    ``field_codes`` isn't supplied.
    """
    paragraphs = list(doc.paragraphs)

    survivors = []
    for para in paragraphs:
        placeholders = _paragraph_placeholders(para)
        if placeholders and all(code in drop_codes for code in placeholders):
            _remove_paragraph(para)
        else:
            survivors.append(para)

    def _is_body(paragraph) -> bool:
        ph = _paragraph_placeholders(paragraph)
        if field_codes:
            return any(c in field_codes and c not in drop_codes for c in ph)
        return bool(paragraph.text.strip())

    for i, para in enumerate(survivors):
        if not _is_heading(para):
            continue
        has_body = False
        for follower in survivors[i + 1:]:
            if _is_heading(follower):
                break
            if _is_body(follower):
                has_body = True
                break
        if not has_body:
            _remove_paragraph(para)


_SECTION_PREFIX_RES = [
    ("cjk", re.compile(r'^\s*([一二三四五六七八九十]+)\s*、\s*(.*)$')),
    ("roman", re.compile(r'^\s*([IVXLCDM]+)\s*\.\s*(.*)$')),
    ("arabic", re.compile(r'^\s*(\d+)\s*[\.、]\s*(.*)$')),
]

_CJK_DIGITS = "零一二三四五六七八九"


def _to_cjk(n: int) -> str:
    """Small int -> CJK numeral (covers 1..99, enough for report sections)."""
    if n < 10:
        return _CJK_DIGITS[n]
    if n < 20:
        return "十" + (_CJK_DIGITS[n - 10] if n > 10 else "")
    if n < 100:
        tens, ones = divmod(n, 10)
        return _CJK_DIGITS[tens] + "十" + (_CJK_DIGITS[ones] if ones else "")
    return str(n)


def _to_roman(n: int) -> str:
    numerals = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    out = ""
    for value, sym in numerals:
        while n >= value:
            out += sym
            n -= value
    return out


def _format_section_prefix(style: str, n: int) -> str:
    if style == "cjk":
        return f"{_to_cjk(n)}、"
    if style == "roman":
        return f"{_to_roman(n)}. "
    return f"{n}. "


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a paragraph's text while preserving its (first run's) formatting."""
    if not paragraph.runs:
        paragraph.text = new_text
        return
    for i, run in enumerate(paragraph.runs):
        run.text = new_text if i == 0 else ""


def _renumber_sections(doc) -> None:
    """Re-sequence numbered section headings so there are no gaps after dropping.

    Only ``Heading`` paragraphs whose text carries a recognizable number prefix
    are touched; the numbering style is taken from the template's own headings
    (CJK for zh, Roman for en) so the output matches the original look. Headings
    without a recognizable prefix are left untouched.
    """
    headings = [p for p in doc.paragraphs
                if (p.style.name or "").startswith("Heading")]
    matched = []
    for p in headings:
        for style, rx in _SECTION_PREFIX_RES:
            m = rx.match(p.text.strip())
            if m:
                matched.append((p, style, m.group(2)))
                break

    styles = [s for _, s, _ in matched]
    if not styles:
        return
    dominant = max(set(styles), key=styles.count)

    for n, (p, _style, title) in enumerate(matched, start=1):
        new_text = _format_section_prefix(dominant, n) + title
        if new_text != p.text.strip():
            _set_paragraph_text(p, new_text)


def _disable_paragraph_numbering(paragraph) -> None:
    """Turn off any (inherited) automatic list numbering on this paragraph.

    Sets an explicit ``numId=0`` override so Word renders no auto number — we
    supply the number as literal text instead, giving deterministic per-section
    restart control.
    """
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:numPr')):
        pPr.remove(existing)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '0')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def _number_list_items(doc) -> None:
    """Prefix each section's list items with a 1-based number that restarts per
    section, and disable inherited auto-numbering so the number isn't doubled.

    The counter resets at every heading, so each subsection's items start at 1.
    Non-list body paragraphs (plain Normal text) are left untouched.
    """
    counter = 0
    for para in doc.paragraphs:
        style = para.style.name or ""
        if style.startswith("Heading") or style == "Title":
            counter = 0
            continue
        if not style.startswith("List"):
            continue
        counter += 1
        text = re.sub(r'^\s*\d+\s*[\.、]\s*', '', para.text.strip())
        _set_paragraph_text(para, f"{counter}. {text}")
        _disable_paragraph_numbering(para)


def _replace_placeholders_in_paragraph(paragraph, field_values: dict):
    """Replace {field_name} placeholders in a paragraph while preserving formatting."""
    full_text = paragraph.text
    if '{' not in full_text:
        return

    placeholders = re.findall(r'\{(\w+)\}', full_text)
    if not placeholders:
        return

    new_text = full_text
    for field_name in placeholders:
        value = field_values.get(field_name, "")
        new_text = new_text.replace(f'{{{field_name}}}', str(value))

    if new_text == full_text:
        return

    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""


def parse_llm_json_response(response_text: str) -> dict:
    """Parse LLM response as JSON, handling common formatting issues."""
    text = response_text.strip()
    if text.startswith('```'):
        text = re.sub(r'^```(?:json)?\s*', '', text)
        text = re.sub(r'\s*```$', '', text)

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        json_match = re.search(r'\{[^{}]*\}', text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except json.JSONDecodeError:
                pass

    logger.warning(f"Failed to parse LLM JSON response: {text[:200]}")
    return {}


